"""
测试向量化脚本 - 用于测试.docx文件的向量化
"""

import os
import docx
import time
from typing import List
from langchain_core.documents import Document
from langchain_community.vectorstores import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter

# 导入我们的嵌入模型
from deepseek_integration import get_embedding_model

def load_docx(file_path):
    """加载Word文档并返回其内容"""
    try:
        doc = docx.Document(file_path)
        text = ""
        
        # 提取文档的所有段落
        for para in doc.paragraphs:
            if para.text.strip():  # 跳过空段落
                text += para.text + "\n"
        
        # 创建Document对象
        metadata = {"source": file_path}
        return Document(page_content=text, metadata=metadata)
    except Exception as e:
        print(f"加载Word文档 {file_path} 时出错: {e}")
        return None

def load_and_vectorize_documents(doc_dir, embedding_model="dashscope", dimensions=1024):
    """加载并向量化文档"""
    # 初始化嵌入模型
    print(f"使用嵌入模型: {embedding_model}, 维度: {dimensions}")
    embedding_function = get_embedding_model(model_name=embedding_model, dimensions=dimensions)
    
    # 设置向量存储
    persist_directory = "chroma_db"
    if not os.path.exists(persist_directory):
        os.makedirs(persist_directory)
    
    # 加载文档
    documents = []
    docx_files = [f for f in os.listdir(doc_dir) if f.endswith('.docx')]
    print(f"找到 {len(docx_files)} 个.docx文件")
    
    for filename in docx_files:
        file_path = os.path.join(doc_dir, filename)
        print(f"加载文档: {file_path}")
        doc = load_docx(file_path)
        if doc:
            print(f"成功加载文档: {filename}, 内容长度: {len(doc.page_content)} 字符")
            documents.append(doc)
    
    print(f"共加载了 {len(documents)} 个文档")
    
    if not documents:
        print("没有成功加载任何文档，向量化过程停止")
        return
    
    # 拆分文档
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=512, chunk_overlap=52)
    all_splits = text_splitter.split_documents(documents)
    print(f"文档拆分完成，共 {len(all_splits)} 个文档块")
    
    # 创建向量存储
    start_time = time.time()
    vectorstore = Chroma.from_documents(
        documents=all_splits,
        embedding=embedding_function,
        persist_directory=persist_directory
    )
    end_time = time.time()
    
    # 持久化存储
    # 注意：自Chroma 0.4.x起，文档会自动持久化，不需要手动调用persist()
    # vectorstore.persist()
    print(f"向量存储已创建并持久化到: {persist_directory}")
    print(f"向量化过程耗时: {end_time - start_time:.2f}秒")
    
    # 测试检索
    print("\n测试检索功能:")
    retriever = vectorstore.as_retriever(
        search_type="similarity",
        search_kwargs={"k": 3}  # 移除score_threshold参数，避免兼容性问题
    )
    
    test_query = "2025年广东省有哪些重要工作计划?"
    print(f"测试查询: '{test_query}'")
    
    docs = retriever.invoke(test_query)
    print(f"检索到 {len(docs)} 个相关文档块")
    
    for i, doc in enumerate(docs):
        print(f"\n文档 {i+1}:")
        print(f"来源: {doc.metadata['source']}")
        print(f"内容预览: {doc.page_content[:150]}...")

if __name__ == "__main__":
    # 文档目录
    doc_dir = "D:\\Cei 需求\\智能体搭建-省区市政府工作报告-守颖"
    
    # 向量化文档，使用支持的维度值
    load_and_vectorize_documents(doc_dir, embedding_model="dashscope", dimensions=1024) 