"""
本地文档加载器 - 用于从本地文件加载文档，特别优化了Word文档处理
"""

import os
import glob
import docx
from typing import List, Optional
from langchain_community.document_loaders import (
    TextLoader,
    PyPDFLoader,
    CSVLoader,
    UnstructuredMarkdownLoader,
    UnstructuredWordDocumentLoader,
)
from langchain_core.documents import Document

def find_files_with_patterns(directory_path: str, patterns: List[str]) -> List[str]:
    """
    在指定目录中查找匹配模式的文件
    
    参数:
        directory_path: 要搜索的目录路径
        patterns: 文件模式列表，例如 ["*.pdf", "*.docx"]
        
    返回:
        匹配的文件路径列表
    """
    all_files = []
    for pattern in patterns:
        pattern_path = os.path.join(directory_path, pattern)
        files = glob.glob(pattern_path)
        all_files.extend(files)
    return all_files

def determine_loader(file_path: str):
    """根据文件扩展名确定合适的加载器"""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == '.txt':
        return TextLoader(file_path)
    elif ext == '.pdf':
        return PyPDFLoader(file_path)
    elif ext == '.csv':
        return CSVLoader(file_path)
    elif ext == '.md':
        return UnstructuredMarkdownLoader(file_path)
    elif ext in ['.doc', '.docx']:
        # 优先使用自定义Word加载器
        try:
            return CustomWordLoader(file_path)
        except Exception as e:
            print(f"使用自定义Word加载器失败: {e}, 尝试使用Unstructured加载器")
            return UnstructuredWordDocumentLoader(file_path)
    else:
        # 默认尝试作为文本加载
        try:
            return TextLoader(file_path)
        except Exception as e:
            print(f"无法加载文件 {file_path}: {e}")
            return None

class CustomWordLoader:
    """
    自定义Word文档加载器，处理更多格式和特殊情况
    """
    
    def __init__(self, file_path: str):
        self.file_path = file_path
        
    def load(self) -> List[Document]:
        """加载Word文档并转换为Document对象"""
        try:
            doc = docx.Document(self.file_path)
            text = ""
            
            # 提取文档的所有段落
            for para in doc.paragraphs:
                if para.text.strip():  # 跳过空段落
                    text += para.text + "\n"
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text for cell in row.cells if cell.text.strip()])
                    if row_text:
                        text += row_text + "\n"
            
            # 创建Document对象
            metadata = {"source": self.file_path}
            return [Document(page_content=text, metadata=metadata)]
            
        except Exception as e:
            print(f"加载Word文档 {self.file_path} 时出错: {e}")
            # 失败时返回空列表
            return []

def load_documents_from_directory(directory_path: str, extensions: Optional[List[str]] = None) -> List[Document]:
    """
    从目录中加载所有符合条件的文档
    
    参数:
        directory_path: 文档所在的目录路径
        extensions: 要加载的文件扩展名列表(例如['.txt', '.pdf'])
        
    返回:
        Document对象列表
    """
    documents = []
    
    print(f"开始扫描目录: {directory_path}")
    print(f"查找文件类型: {extensions if extensions else '所有文件'}")
    
    for root, _, files in os.walk(directory_path):
        for file in files:
            file_path = os.path.join(root, file)
            file_ext = os.path.splitext(file)[1].lower()
            
            # 如果指定了扩展名，则只加载匹配的文件
            if extensions and file_ext not in extensions:
                continue
                
            loader = determine_loader(file_path)
            if loader:
                try:
                    docs = loader.load()
                    documents.extend(docs)
                    print(f"已加载文件: {file_path}")
                except Exception as e:
                    print(f"加载文件 {file_path} 时出错: {e}")
    
    print(f"共加载了 {len(documents)} 个文档")
    return documents 